"""Export API Routes.

Provides endpoints for exporting analysis results to various formats.
"""

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import io
from datetime import datetime
import sys
from pathlib import Path

# Add src to path
src_path = Path(__file__).parent.parent.parent.parent / "src"
sys.path.insert(0, str(src_path))

from services import session_manager

router = APIRouter(prefix="/export", tags=["export"])


class ExportOptions(BaseModel):
    includeMethods: bool = True
    includeResults: bool = True
    includeFigures: bool = True
    includeTables: bool = True
    includeCode: bool = False
    includeInterpretations: bool = True
    journalStyle: str = "nejm"


class ExportRequest(BaseModel):
    format: str
    options: ExportOptions


@router.post("/report/{session_id}")
async def export_report(session_id: str, request: ExportRequest):
    """Export analysis report to specified format."""
    session = session_manager.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if not session.analyses:
        raise HTTPException(status_code=400, detail="No analyses to export")

    try:
        if request.format == "docx":
            content, media_type, filename = _export_docx(session, request.options)
        elif request.format == "pdf":
            content, media_type, filename = _export_pdf(session, request.options)
        elif request.format == "html":
            content, media_type, filename = _export_html(session, request.options)
        elif request.format == "latex":
            content, media_type, filename = _export_latex(session, request.options)
        elif request.format == "markdown":
            content, media_type, filename = _export_markdown(session, request.options)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported format: {request.format}")

        return StreamingResponse(
            io.BytesIO(content),
            media_type=media_type,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/tables/{session_id}")
async def export_tables(session_id: str, format: str = "xlsx"):
    """Export all analysis tables."""
    session = session_manager.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    try:
        if format == "xlsx":
            content = _export_tables_xlsx(session)
            media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            filename = "analysis_tables.xlsx"
        elif format == "csv":
            content = _export_tables_csv(session)
            media_type = "application/zip"
            filename = "analysis_tables.zip"
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported format: {format}")

        return StreamingResponse(
            io.BytesIO(content),
            media_type=media_type,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def _export_docx(session, options: ExportOptions) -> tuple:
    """Export to Word document."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed")

    doc = Document()

    # Title
    title = doc.add_heading("Statistical Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    if session.data is not None:
        doc.add_paragraph(f"Dataset: {session.data_profile.get('filename', 'Unknown')}")
        doc.add_paragraph(f"Sample size: {len(session.data):,} observations")

    doc.add_paragraph()

    # Methods Section
    if options.includeMethods:
        doc.add_heading("Methods", level=1)
        methods_text = _generate_methods_section(session, options.journalStyle)
        doc.add_paragraph(methods_text)

    # Results Section
    if options.includeResults:
        doc.add_heading("Results", level=1)
        for analysis in session.analyses:
            doc.add_heading(analysis.get('analysis_type', 'Analysis').replace('_', ' ').title(), level=2)

            # Add interpretation if available
            if options.includeInterpretations and 'interpretation' in analysis:
                doc.add_paragraph(analysis['interpretation'])

            # Add key statistics
            results = analysis.get('results', {})
            if 'p_value' in results:
                doc.add_paragraph(f"p-value: {results['p_value']:.4f}")

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "analysis_report.docx"


def _export_pdf(session, options: ExportOptions) -> tuple:
    """Export to PDF."""
    try:
        from fpdf import FPDF
    except ImportError:
        raise HTTPException(status_code=500, detail="fpdf2 not installed")

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)

    # Title
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Statistical Analysis Report", ln=True, align="C")
    pdf.ln(10)

    # Metadata
    pdf.set_font("Helvetica", size=10)
    pdf.cell(0, 6, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=True)

    # Methods
    if options.includeMethods:
        pdf.ln(5)
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 10, "Methods", ln=True)
        pdf.set_font("Helvetica", size=10)
        methods_text = _generate_methods_section(session, options.journalStyle)
        pdf.multi_cell(0, 5, methods_text)

    # Results
    if options.includeResults:
        pdf.ln(5)
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 10, "Results", ln=True)

        for analysis in session.analyses:
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, analysis.get('analysis_type', 'Analysis').replace('_', ' ').title(), ln=True)

            pdf.set_font("Helvetica", size=10)
            if options.includeInterpretations and 'interpretation' in analysis:
                pdf.multi_cell(0, 5, analysis['interpretation'][:500])  # Truncate long text

            results = analysis.get('results', {})
            if 'p_value' in results:
                pdf.cell(0, 5, f"p-value: {results['p_value']:.4f}", ln=True)
            pdf.ln(3)

    buffer = io.BytesIO()
    pdf.output(buffer)
    buffer.seek(0)

    return buffer.read(), "application/pdf", "analysis_report.pdf"


def _export_html(session, options: ExportOptions) -> tuple:
    """Export to HTML."""
    html_parts = [
        "<!DOCTYPE html>",
        "<html>",
        "<head>",
        "<title>Statistical Analysis Report</title>",
        "<style>",
        "body { font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }",
        "h1 { color: #333; border-bottom: 2px solid #4F46E5; padding-bottom: 10px; }",
        "h2 { color: #4F46E5; }",
        "table { border-collapse: collapse; width: 100%; margin: 20px 0; }",
        "th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }",
        "th { background-color: #4F46E5; color: white; }",
        ".stat { background-color: #f8f9fa; padding: 10px; border-radius: 5px; margin: 10px 0; }",
        ".significant { color: #10B981; font-weight: bold; }",
        "</style>",
        "</head>",
        "<body>",
        f"<h1>Statistical Analysis Report</h1>",
        f"<p>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>",
    ]

    if options.includeMethods:
        html_parts.append("<h2>Methods</h2>")
        methods_text = _generate_methods_section(session, options.journalStyle)
        html_parts.append(f"<p>{methods_text}</p>")

    if options.includeResults:
        html_parts.append("<h2>Results</h2>")
        for analysis in session.analyses:
            html_parts.append(f"<h3>{analysis.get('analysis_type', 'Analysis').replace('_', ' ').title()}</h3>")

            if options.includeInterpretations and 'interpretation' in analysis:
                html_parts.append(f"<p>{analysis['interpretation']}</p>")

            results = analysis.get('results', {})
            if 'p_value' in results:
                sig_class = 'significant' if results['p_value'] < 0.05 else ''
                html_parts.append(f'<p class="stat {sig_class}">p-value: {results["p_value"]:.4f}</p>')

    html_parts.extend(["</body>", "</html>"])

    content = "\n".join(html_parts).encode('utf-8')
    return content, "text/html", "analysis_report.html"


def _export_latex(session, options: ExportOptions) -> tuple:
    """Export to LaTeX."""
    lines = [
        r"\documentclass{article}",
        r"\usepackage{booktabs}",
        r"\usepackage{graphicx}",
        r"\usepackage{hyperref}",
        "",
        r"\begin{document}",
        "",
        r"\title{Statistical Analysis Report}",
        r"\maketitle",
        "",
    ]

    if options.includeMethods:
        lines.append(r"\section{Methods}")
        methods_text = _generate_methods_section(session, options.journalStyle)
        lines.append(_escape_latex(methods_text))
        lines.append("")

    if options.includeResults:
        lines.append(r"\section{Results}")
        for analysis in session.analyses:
            lines.append(r"\subsection{" + _escape_latex(analysis.get('analysis_type', 'Analysis').replace('_', ' ').title()) + "}")

            if options.includeInterpretations and 'interpretation' in analysis:
                lines.append(_escape_latex(analysis['interpretation']))

            results = analysis.get('results', {})
            if 'p_value' in results:
                lines.append(f"p-value: {results['p_value']:.4f}")
            lines.append("")

    lines.append(r"\end{document}")

    content = "\n".join(lines).encode('utf-8')
    return content, "application/x-latex", "analysis_report.tex"


def _export_markdown(session, options: ExportOptions) -> tuple:
    """Export to Markdown."""
    lines = [
        "# Statistical Analysis Report",
        "",
        f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}*",
        "",
    ]

    if options.includeMethods:
        lines.append("## Methods")
        lines.append("")
        methods_text = _generate_methods_section(session, options.journalStyle)
        lines.append(methods_text)
        lines.append("")

    if options.includeResults:
        lines.append("## Results")
        lines.append("")
        for analysis in session.analyses:
            lines.append(f"### {analysis.get('analysis_type', 'Analysis').replace('_', ' ').title()}")
            lines.append("")

            if options.includeInterpretations and 'interpretation' in analysis:
                lines.append(analysis['interpretation'])
                lines.append("")

            results = analysis.get('results', {})
            if 'p_value' in results:
                sig = "**" if results['p_value'] < 0.05 else ""
                lines.append(f"- p-value: {sig}{results['p_value']:.4f}{sig}")
            lines.append("")

    content = "\n".join(lines).encode('utf-8')
    return content, "text/markdown", "analysis_report.md"


def _export_tables_xlsx(session) -> bytes:
    """Export all tables to Excel."""
    import pandas as pd

    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        for i, analysis in enumerate(session.analyses):
            results = analysis.get('results', {})

            # Try to extract tabular data
            if 'table_data' in results:
                df = pd.DataFrame(results['table_data'])
                sheet_name = f"{analysis.get('analysis_type', 'Analysis')[:28]}_{i+1}"
                df.to_excel(writer, sheet_name=sheet_name, index=False)
            elif 'coefficients' in results:
                df = pd.DataFrame(results['coefficients']).T
                sheet_name = f"Coefficients_{i+1}"
                df.to_excel(writer, sheet_name=sheet_name)

    buffer.seek(0)
    return buffer.read()


def _export_tables_csv(session) -> bytes:
    """Export all tables to CSV (zipped)."""
    import zipfile
    import pandas as pd

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for i, analysis in enumerate(session.analyses):
            results = analysis.get('results', {})

            if 'table_data' in results:
                df = pd.DataFrame(results['table_data'])
                csv_buffer = io.StringIO()
                df.to_csv(csv_buffer, index=False)
                filename = f"{analysis.get('analysis_type', 'analysis')}_{i+1}.csv"
                zf.writestr(filename, csv_buffer.getvalue())

    buffer.seek(0)
    return buffer.read()


def _generate_methods_section(session, journal_style: str) -> str:
    """Generate methods section text."""
    analyses = session.analyses
    if not analyses:
        return "No analyses performed."

    analysis_types = set(a.get('analysis_type', '') for a in analyses)

    methods = []
    methods.append("Statistical analyses were performed using Python-based statistical software.")

    if session.data is not None:
        methods.append(f"The analysis included {len(session.data):,} observations.")

    # Describe analyses
    if 'ttest' in analysis_types or 't_test' in analysis_types:
        methods.append("Continuous variables were compared between groups using Student's t-test or Welch's t-test for unequal variances.")

    if 'chi_square' in analysis_types:
        methods.append("Categorical variables were compared using chi-square tests or Fisher's exact test when expected cell counts were less than 5.")

    if 'logistic_regression' in analysis_types:
        methods.append("Multivariable logistic regression was used to estimate adjusted odds ratios with 95% confidence intervals.")

    if 'cox_regression' in analysis_types:
        methods.append("Cox proportional hazards regression was used to estimate hazard ratios. The proportional hazards assumption was assessed using Schoenfeld residuals.")

    if 'kaplan_meier' in analysis_types:
        methods.append("Survival curves were estimated using the Kaplan-Meier method and compared using the log-rank test.")

    methods.append("Two-sided p-values < 0.05 were considered statistically significant.")

    return " ".join(methods)


def _escape_latex(text: str) -> str:
    """Escape special LaTeX characters."""
    replacements = {
        '&': r'\&', '%': r'\%', '$': r'\$', '#': r'\#',
        '_': r'\_', '{': r'\{', '}': r'\}',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text
