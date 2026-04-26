"""Chat and LLM interaction endpoints."""

from fastapi import APIRouter, HTTPException, WebSocket, WebSocketDisconnect
from fastapi.responses import JSONResponse
from typing import Dict, Any, List, Optional
from pydantic import BaseModel
from datetime import datetime
import json
import math
import pandas as pd


class PlanModifications(BaseModel):
    """Modifications to an analysis plan."""
    variable_mappings: Optional[Dict[str, str]] = None
    exclude_analyses: Optional[List[str]] = None
    additional_covariates: Optional[List[str]] = None


def sanitize_for_json(obj):
    """Recursively sanitize an object for JSON serialization.

    Converts NaN and Infinity to None, and numpy types to Python types.
    """
    import numpy as np

    if isinstance(obj, dict):
        return {k: sanitize_for_json(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [sanitize_for_json(item) for item in obj]
    elif isinstance(obj, (np.bool_,)):
        return bool(obj)
    elif isinstance(obj, np.integer):
        return int(obj)
    elif isinstance(obj, np.floating):
        if np.isnan(obj) or np.isinf(obj):
            return None
        return float(obj)
    elif isinstance(obj, np.ndarray):
        return sanitize_for_json(obj.tolist())
    elif isinstance(obj, float):
        if math.isnan(obj) or math.isinf(obj):
            return None
        return obj
    elif isinstance(obj, bool):
        return obj
    else:
        return obj

from services import session_manager, llm_service, analysis_service
from models import (
    ChatRequest, ChatResponse,
    ResearchQuestionRequest, ResearchQuestionResponse
)

router = APIRouter(prefix="/chat", tags=["Chat"])


@router.post("/message")
async def send_message(request: ChatRequest):
    """Send a chat message and get a response.

    The assistant can answer questions about statistics,
    recommend analyses, and help interpret results.
    """
    session = session_manager.get_session(request.session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    # Build context
    context = {
        "data_loaded": session.data is not None,
        "filename": session.filename,
        "n_rows": len(session.data) if session.data is not None else 0,
        "n_cols": len(session.data.columns) if session.data is not None else 0,
        "variables": list(session.data.columns)[:20] if session.data is not None else [],
        "analyses_performed": [a.get("analysis_type") for a in session.analyses]
    }

    # Add any additional context from request
    if request.context:
        context.update(request.context)

    # Get chat response
    response = llm_service.chat(
        request.message,
        context,
        session.chat_history
    )

    # Store messages
    session_manager.add_chat_message(request.session_id, "user", request.message)
    session_manager.add_chat_message(request.session_id, "assistant", response["message"])

    # Check if there's an action to execute
    action_result = None
    if response.get("action"):
        action = response["action"]
        if action.get("action") == "run_analysis" and session.data is not None:
            action_result = analysis_service.run_analysis(
                session.data,
                action.get("analysis_type"),
                action.get("parameters", {})
            )
            if action_result.get("success"):
                session_manager.add_analysis(request.session_id, action_result)

    return JSONResponse(content={
        "success": True,
        "message": response["message"],
        "action_executed": action_result is not None,
        "action_result": action_result
    })


@router.post("/research-question")
async def analyze_research_question(request: ResearchQuestionRequest):
    """Analyze a research question and get recommended analyses.

    Returns a comprehensive analysis plan based on the research
    question and the loaded data.
    """
    session = session_manager.get_session(request.session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if session.data is None or session.data_profile is None:
        raise HTTPException(
            status_code=400,
            detail="Please upload data before asking a research question"
        )

    # Get analysis plan from LLM
    result = llm_service.analyze_research_question(
        request.question,
        session.data_profile
    )

    if result.get("success"):
        session.analysis_plan = result["plan"]

        # Store in chat history
        session_manager.add_chat_message(
            request.session_id, "user",
            f"Research question: {request.question}"
        )
        session_manager.add_chat_message(
            request.session_id, "assistant",
            f"I've created an analysis plan with {len(result['plan'].get('primary_analyses', []))} "
            f"primary analyses and {len(result['plan'].get('secondary_analyses', []))} secondary analyses."
        )

    return JSONResponse(content=result)


@router.post("/confirm-plan/{session_id}")
async def confirm_analysis_plan(
    session_id: str,
    modifications: Optional[PlanModifications] = None
):
    """Confirm the analysis plan before execution.

    Optionally accept modifications to variable mappings.
    Returns the confirmed plan ready for execution.
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if not session.analysis_plan:
        raise HTTPException(status_code=404, detail="No analysis plan found. Please ask a research question first.")

    plan = session.analysis_plan

    # Apply modifications if provided
    if modifications:
        if modifications.variable_mappings:
            _apply_variable_mappings(plan, modifications.variable_mappings)

        if modifications.exclude_analyses:
            plan["primary_analyses"] = [
                a for a in plan.get("primary_analyses", [])
                if a.get("test_name") not in modifications.exclude_analyses
            ]
            plan["secondary_analyses"] = [
                a for a in plan.get("secondary_analyses", [])
                if a.get("test_name") not in modifications.exclude_analyses
            ]

        if modifications.additional_covariates:
            _add_covariates_to_plan(plan, modifications.additional_covariates)

    # Mark as confirmed
    plan["confirmed"] = True
    plan["confirmed_at"] = datetime.now().isoformat()

    # Store in chat history
    session_manager.add_chat_message(
        session_id, "assistant",
        f"Analysis plan confirmed with {len(plan.get('primary_analyses', []))} primary analyses. Ready for execution."
    )

    return JSONResponse(content={
        "success": True,
        "message": "Plan confirmed. Ready for execution.",
        "plan": sanitize_for_json(plan)
    })


def _apply_variable_mappings(plan: Dict, mappings: Dict[str, str]) -> None:
    """Apply user-provided variable mappings to the plan."""
    for analysis in plan.get("primary_analyses", []) + plan.get("secondary_analyses", []):
        api_call = analysis.get("api_call", {})
        params = api_call.get("parameters", {})

        for key, value in params.items():
            if isinstance(value, str) and value in mappings:
                params[key] = mappings[value]
            elif isinstance(value, list):
                params[key] = [mappings.get(v, v) if isinstance(v, str) else v for v in value]


def _add_covariates_to_plan(plan: Dict, additional_covariates: List[str]) -> None:
    """Add additional covariates to regression analyses."""
    regression_types = ["logistic_regression", "linear_regression", "cox_regression"]

    for analysis in plan.get("primary_analyses", []) + plan.get("secondary_analyses", []):
        api_call = analysis.get("api_call", {})
        if api_call.get("analysis_type") in regression_types:
            params = api_call.get("parameters", {})

            # Add to predictors or covariates
            for param_name in ["predictors", "covariates"]:
                if param_name in params:
                    existing = params[param_name] if isinstance(params[param_name], list) else []
                    for cov in additional_covariates:
                        if cov not in existing:
                            existing.append(cov)
                    params[param_name] = existing
                    break
            else:
                # If no predictors/covariates field exists, add one
                params["covariates"] = additional_covariates


@router.get("/plan/{session_id}")
async def get_analysis_plan(session_id: str):
    """Get the current analysis plan for a session."""
    session = session_manager.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if not session.analysis_plan:
        return JSONResponse(content={
            "success": False,
            "message": "No analysis plan found"
        })

    return JSONResponse(content={
        "success": True,
        "plan": sanitize_for_json(session.analysis_plan),
        "confirmed": session.analysis_plan.get("confirmed", False)
    })


@router.get("/history/{session_id}")
async def get_chat_history(session_id: str):
    """Get chat history for a session."""
    history = session_manager.get_chat_history(session_id)
    return JSONResponse(content={
        "success": True,
        "history": history
    })


def _normalize_analysis_params(analysis_type: str, params: Dict[str, Any]) -> tuple:
    """Normalize LLM-generated parameters to match analysis service expectations."""
    # Map analysis type variations
    type_mapping = {
        "anova": "one_way_anova",
        "t_test": "independent_ttest",
        "ttest": "independent_ttest",
        "survival": "kaplan_meier",
        "km": "kaplan_meier",
        "chi_square_test": "chi_square",
        "chisquare": "chi_square",
        "chi-square": "chi_square",
        "fisher": "fisher_exact",
        "fishers_exact": "fisher_exact",
        "cox": "cox_regression",
        "cox_ph": "cox_regression",
        "cox_proportional_hazards": "cox_regression",
        "logistic": "logistic_regression",
        "linear": "linear_regression",
        "multiple_regression": "linear_regression",
        "multivariable_regression": "linear_regression",
        "pearson": "pearson_correlation",
        "spearman": "spearman_correlation",
    }
    normalized_type = type_mapping.get(analysis_type, analysis_type)

    # Map parameter name variations
    param_mapping = {
        "time_var": "time",
        "event_var": "event",
        "group_var": "group",
        "outcome_var": "outcome",
        "predictor_var": "predictors",
        "covariate_vars": "covariates",
        "row_variable": "row_var",
        "col_variable": "col_var",
        "variable_1": "var1",
        "variable_2": "var2",
        "x_variable": "var1",
        "y_variable": "var2",
    }

    normalized_params = {}
    for key, value in params.items():
        new_key = param_mapping.get(key, key)
        normalized_params[new_key] = value

    # Handle chi-square which expects row_var and col_var
    if normalized_type == "chi_square":
        if "outcome" in normalized_params and "group" in normalized_params:
            normalized_params["row_var"] = normalized_params.pop("outcome")
            normalized_params["col_var"] = normalized_params.pop("group")

    # Handle logistic/linear regression
    if normalized_type in ["logistic_regression", "linear_regression"]:
        if "outcome" in normalized_params:
            normalized_params["outcome"] = normalized_params["outcome"]
        if "covariates" in normalized_params and "group" in normalized_params:
            # Add group to predictors
            covariates = normalized_params.get("covariates", [])
            group = normalized_params.pop("group", None)
            if group and group not in covariates:
                covariates = [group] + list(covariates)
            normalized_params["predictors"] = covariates

    # Handle t-test and ANOVA
    if normalized_type in ["independent_ttest", "one_way_anova", "mann_whitney", "kruskal_wallis"]:
        if "outcome" in normalized_params:
            normalized_params["outcome"] = normalized_params["outcome"]
        if "group" in normalized_params:
            normalized_params["group"] = normalized_params["group"]

    return normalized_type, normalized_params


class ExecutePlanRequest(BaseModel):
    """Request body for execute-plan endpoint."""
    analyses: Optional[List[str]] = None


@router.post("/execute-plan/{session_id}")
async def execute_analysis_plan(session_id: str, request: Optional[ExecutePlanRequest] = None):
    """Execute analyses from the analysis plan.

    If analyses list is provided, only those are executed.
    Otherwise, all primary and secondary analyses are executed.
    """
    analyses = request.analyses if request else None
    session = session_manager.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if not session.analysis_plan:
        raise HTTPException(status_code=400, detail="No analysis plan found")

    if session.data is None:
        raise HTTPException(status_code=400, detail="No data loaded")

    plan = session.analysis_plan

    # Check if plan requires confirmation
    if plan.get("require_confirmation") and not plan.get("confirmed"):
        return JSONResponse(content={
            "success": False,
            "error": "Plan requires confirmation before execution",
            "require_confirmation": True,
            "plan": sanitize_for_json(plan)
        }, status_code=400)

    results = []
    errors = []

    # Get analyses to run (both primary and secondary)
    all_analyses = plan.get("primary_analyses", []) + plan.get("secondary_analyses", [])
    if analyses:
        analyses_to_run = [a for a in all_analyses if a.get("test_name") in analyses]
    else:
        analyses_to_run = all_analyses

    for analysis in analyses_to_run:
        api_call = analysis.get("api_call", {})
        if api_call:
            # Normalize parameters for the analysis service
            analysis_type, params = _normalize_analysis_params(
                api_call.get("analysis_type", ""),
                api_call.get("parameters", {})
            )

            try:
                result = analysis_service.run_analysis(
                    session.data,
                    analysis_type,
                    params
                )

                # Add interpretation and metadata to successful results
                if result.get("success"):
                    result = _add_interpretation(result, analysis_type, analysis.get("rationale", ""))
                    # Add heading info to result for display
                    result["test_name"] = analysis.get("test_name")
                    result["rationale"] = analysis.get("rationale")
                    session_manager.add_analysis(session_id, result)

                results.append({
                    "test_name": analysis.get("test_name"),
                    "rationale": analysis.get("rationale"),
                    "result": sanitize_for_json(result)
                })
            except Exception as e:
                errors.append({
                    "test_name": analysis.get("test_name"),
                    "error": str(e)
                })

    # Generate summary message
    n_success = sum(1 for r in results if r.get("result", {}).get("success"))
    summary_message = f"Executed {n_success} of {len(analyses_to_run)} analyses successfully."
    if errors:
        summary_message += f" {len(errors)} analyses failed."

    # Store summary in chat
    session_manager.add_chat_message(session_id, "assistant", summary_message)

    return JSONResponse(content=sanitize_for_json({
        "success": len(errors) == 0,
        "message": summary_message,
        "n_executed": n_success,
        "n_total": len(analyses_to_run),
        "results": results,
        "errors": errors if errors else None
    }))


def _add_interpretation(result: Dict[str, Any], analysis_type: str,
                       rationale: str = "") -> Dict[str, Any]:
    """Add clinical interpretation to analysis results."""
    results_data = result.get("results", {})

    interpretation = ""

    if analysis_type == "chi_square":
        p = results_data.get("p_value")
        stat = results_data.get("statistic")
        cramers_v = results_data.get("cramers_v") or results_data.get("effect_size")

        # Handle None or non-numeric values
        try:
            p_val = float(p) if p is not None else 1.0
            stat_val = float(stat) if stat is not None else 0.0
        except (TypeError, ValueError):
            p_val = 1.0
            stat_val = 0.0

        sig = "statistically significant" if p_val < 0.05 else "not statistically significant"
        p_str = f"{p_val:.4f}" if p_val >= 0.001 else "<0.001"

        interpretation = f"""**Chi-Square Test Results**

- Chi-square statistic: {stat_val:.3f}
- P-value: {p_str}
- Conclusion: The association is **{sig}** (α = 0.05)
"""
        if cramers_v is not None:
            try:
                cv = float(cramers_v)
                effect_label = "negligible" if cv < 0.1 else "small" if cv < 0.3 else "medium" if cv < 0.5 else "large"
                interpretation += f"- Effect size (Cramér's V): {cv:.3f} ({effect_label})\n"
            except (TypeError, ValueError):
                pass

    elif analysis_type == "logistic_regression":
        n_obs = results_data.get("n_observations", 0) or 0
        pseudo_r2 = results_data.get("pseudo_r2")
        auc = results_data.get("auc")
        coefficients = results_data.get("coefficients", {})

        # Handle None values
        try:
            r2_val = float(pseudo_r2) if pseudo_r2 is not None else 0.0
            auc_val = float(auc) if auc is not None else 0.0
        except (TypeError, ValueError):
            r2_val = 0.0
            auc_val = 0.0

        auc_label = "excellent" if auc_val >= 0.9 else "good" if auc_val >= 0.8 else "acceptable" if auc_val >= 0.7 else "poor"

        interpretation = f"""**Logistic Regression Results**

- Observations: {n_obs}
- Pseudo R²: {r2_val:.3f}
- AUC: {auc_val:.3f} ({auc_label} discrimination)

**Odds Ratios (95% CI):**
"""
        for var_name, stats in coefficients.items():
            if var_name.lower() != "intercept" and var_name.lower() != "const":
                try:
                    coef = stats.get("coefficient", 0) if stats.get("coefficient") is not None else 0
                    odds_ratio = stats.get("odds_ratio") or (2.718 ** coef)
                    odds_ratio = float(odds_ratio) if odds_ratio is not None else 1.0
                    ci_lower = float(stats.get("ci_lower", 0)) if stats.get("ci_lower") is not None else 0.0
                    ci_upper = float(stats.get("ci_upper", 0)) if stats.get("ci_upper") is not None else 0.0
                    p_val = float(stats.get("p_value", 1)) if stats.get("p_value") is not None else 1.0
                    sig_marker = "*" if p_val < 0.05 else ""
                    p_str = f"{p_val:.3f}" if p_val >= 0.001 else "<0.001"
                    interpretation += f"- {var_name}: OR={odds_ratio:.2f} ({ci_lower:.2f}-{ci_upper:.2f}), p={p_str}{sig_marker}\n"
                except (TypeError, ValueError):
                    interpretation += f"- {var_name}: (calculation error)\n"

        interpretation += """
**Interpretation:**
- OR > 1: Higher odds of outcome
- OR < 1: Lower odds of outcome
- CI not crossing 1: Statistically significant
"""

    elif analysis_type == "cox_regression":
        n_obs = results_data.get("n_observations", 0) or 0
        n_events = results_data.get("n_events", 0) or 0
        coefficients = results_data.get("coefficients", {})

        interpretation = f"""**Cox Proportional Hazards Regression**

- Observations: {n_obs}
- Events: {n_events}
- Censored: {n_obs - n_events}

**Hazard Ratios (95% CI):**
"""
        for var_name, stats in coefficients.items():
            try:
                hr = float(stats.get("hazard_ratio", 1)) if stats.get("hazard_ratio") is not None else 1.0
                ci_lower = float(stats.get("hr_ci_lower", 0)) if stats.get("hr_ci_lower") is not None else 0.0
                ci_upper = float(stats.get("hr_ci_upper", 0)) if stats.get("hr_ci_upper") is not None else 0.0
                p_val = float(stats.get("p_value", 1)) if stats.get("p_value") is not None else 1.0
                sig_marker = "*" if p_val < 0.05 else ""
                p_str = f"{p_val:.3f}" if p_val >= 0.001 else "<0.001"
                if hr != 1:
                    effect = f"{abs((hr - 1) * 100):.1f}% {'increased' if hr > 1 else 'decreased'} risk"
                else:
                    effect = "no effect"
                interpretation += f"- {var_name}: HR={hr:.2f} ({ci_lower:.2f}-{ci_upper:.2f}), p={p_str}{sig_marker} ({effect})\n"
            except (TypeError, ValueError):
                interpretation += f"- {var_name}: (calculation error)\n"

        interpretation += """
**Interpretation:**
- HR > 1: Higher hazard (worse survival)
- HR < 1: Lower hazard (better survival)
- CI not crossing 1: Statistically significant
"""

    elif analysis_type in ["kaplan_meier", "log_rank"]:
        p = results_data.get("p_value")
        if p is not None:
            try:
                p_val = float(p)
                sig = "significant difference" if p_val < 0.05 else "no significant difference"
                p_str = f"{p_val:.4f}" if p_val >= 0.001 else "<0.001"
                interpretation = f"""**Kaplan-Meier Survival Analysis**

- Log-rank p-value: {p_str}
- Conclusion: There is **{sig}** in survival between groups (α = 0.05)
"""
            except (TypeError, ValueError):
                interpretation = "**Kaplan-Meier Survival Analysis** - Results computed."

    elif analysis_type == "one_way_anova":
        p = results_data.get("p_value")
        f_stat = results_data.get("statistic")
        eta_sq = results_data.get("eta_squared")

        try:
            p_val = float(p) if p is not None else 1.0
            f_val = float(f_stat) if f_stat is not None else 0.0
        except (TypeError, ValueError):
            p_val = 1.0
            f_val = 0.0

        sig = "significant" if p_val < 0.05 else "not significant"
        p_str = f"{p_val:.4f}" if p_val >= 0.001 else "<0.001"

        interpretation = f"""**One-Way ANOVA Results**

- F-statistic: {f_val:.3f}
- P-value: {p_str}
- Conclusion: Group differences are **{sig}** (α = 0.05)
"""
        if eta_sq is not None:
            try:
                eta = float(eta_sq)
                effect_label = "small" if eta < 0.06 else "medium" if eta < 0.14 else "large"
                interpretation += f"- Effect size (η²): {eta:.3f} ({effect_label})\n"
            except (TypeError, ValueError):
                pass

    elif analysis_type == "table1":
        interpretation = """**Table 1: Baseline Characteristics**

This table presents the baseline characteristics of the study population stratified by group.
- Continuous variables are presented as mean ± SD or median [IQR]
- Categorical variables are presented as n (%)
- P-values compare differences between groups
"""

    elif analysis_type in ["independent_ttest", "paired_ttest"]:
        p = results_data.get("p_value")
        t_stat = results_data.get("statistic")
        cohens_d = results_data.get("cohens_d") or results_data.get("effect_size")

        try:
            p_val = float(p) if p is not None else 1.0
            t_val = float(t_stat) if t_stat is not None else 0.0
        except (TypeError, ValueError):
            p_val = 1.0
            t_val = 0.0

        sig = "significant" if p_val < 0.05 else "not significant"
        p_str = f"{p_val:.4f}" if p_val >= 0.001 else "<0.001"

        interpretation = f"""**T-Test Results**

- t-statistic: {t_val:.3f}
- P-value: {p_str}
- Conclusion: The difference is **{sig}** (α = 0.05)
"""
        if cohens_d is not None:
            try:
                d = float(cohens_d)
                effect_label = "small" if abs(d) < 0.5 else "medium" if abs(d) < 0.8 else "large"
                interpretation += f"- Effect size (Cohen's d): {d:.3f} ({effect_label})\n"
            except (TypeError, ValueError):
                pass

    # Add rationale context if provided
    if rationale and interpretation:
        interpretation = f"**Purpose:** {rationale}\n\n{interpretation}"

    result["interpretation"] = interpretation
    return result


@router.get("/confounders/{session_id}")
async def get_confounder_suggestions(
    session_id: str,
    outcome: str = None,
    outcome_type: str = "binary",  # binary, continuous, survival
    time_col: str = None,
    event_col: str = None
):
    """Get ALL variables with univariate significance testing.

    Returns ALL variables from the dataset (excluding outcome/time/event columns)
    with their univariate p-values to help users decide which to include
    in adjusted analyses.
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if session.data is None:
        raise HTTPException(status_code=400, detail="No data loaded")

    df = session.data

    # Build list of columns to exclude (outcome, time, event columns)
    exclude_cols = set()
    if outcome:
        exclude_cols.add(outcome)
    if time_col:
        exclude_cols.add(time_col)
    if event_col:
        exclude_cols.add(event_col)

    # Get ALL columns from the dataset except excluded ones
    confounders = [col for col in df.columns if col not in exclude_cols]

    results = []

    for conf in confounders:
        conf_info = {
            "variable": conf,
            "n_valid": int(df[conf].notna().sum()),
            "missing_pct": float(df[conf].isna().mean() * 100),
            "type": "continuous" if df[conf].dtype in ['float64', 'int64'] and df[conf].nunique() > 10 else "categorical"
        }

        # Run univariate test if outcome is provided
        if outcome and outcome in df.columns:
            try:
                if outcome_type == "binary":
                    # Chi-square for categorical, t-test for continuous
                    if conf_info["type"] == "categorical":
                        from scipy import stats
                        contingency = pd.crosstab(df[conf], df[outcome])
                        chi2, p_val, _, _ = stats.chi2_contingency(contingency)
                        conf_info["test"] = "chi-square"
                        conf_info["statistic"] = float(chi2)
                        conf_info["p_value"] = float(p_val)
                    else:
                        from scipy import stats
                        group0 = df[df[outcome] == 0][conf].dropna()
                        group1 = df[df[outcome] == 1][conf].dropna()
                        if len(group0) > 1 and len(group1) > 1:
                            stat, p_val = stats.ttest_ind(group0, group1)
                            conf_info["test"] = "t-test"
                            conf_info["statistic"] = float(stat)
                            conf_info["p_value"] = float(p_val)
                elif outcome_type == "continuous":
                    from scipy import stats
                    if conf_info["type"] == "categorical":
                        # ANOVA
                        groups = [df[df[conf] == cat][outcome].dropna() for cat in df[conf].unique() if pd.notna(cat)]
                        groups = [g for g in groups if len(g) > 0]
                        if len(groups) >= 2:
                            stat, p_val = stats.f_oneway(*groups)
                            conf_info["test"] = "ANOVA"
                            conf_info["statistic"] = float(stat)
                            conf_info["p_value"] = float(p_val)
                    else:
                        # Correlation
                        clean = df[[conf, outcome]].dropna()
                        if len(clean) > 2:
                            r, p_val = stats.pearsonr(clean[conf], clean[outcome])
                            conf_info["test"] = "correlation"
                            conf_info["statistic"] = float(r)
                            conf_info["p_value"] = float(p_val)

                conf_info["significant"] = conf_info.get("p_value", 1) < 0.05
            except Exception as e:
                conf_info["test_error"] = str(e)

        results.append(conf_info)

    # Sort by p-value (significant first)
    results.sort(key=lambda x: (not x.get("significant", False), x.get("p_value", 1)))

    return JSONResponse(content={
        "success": True,
        "outcome": outcome,
        "outcome_type": outcome_type,
        "confounders": results,
        "all_columns": list(df.columns)
    })


@router.post("/run-adjusted-analysis/{session_id}")
async def run_adjusted_analysis(session_id: str, request: Dict[str, Any]):
    """Run adjusted analysis with user-selected confounders.

    Request body:
    {
        "analysis_type": "logistic_regression" | "cox_regression",
        "outcome": "column_name",
        "predictors": ["var1", "var2"],  # Main exposure variables
        "confounders": ["age", "sex", ...],  # User-selected confounders
        "time": "time_col",  # For Cox only
        "event": "event_col"  # For Cox only
    }
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if session.data is None:
        raise HTTPException(status_code=400, detail="No data loaded")

    analysis_type = request.get("analysis_type")
    outcome = request.get("outcome")
    predictors = request.get("predictors", [])
    confounders = request.get("confounders", [])

    # Combine predictors and confounders
    all_predictors = list(predictors) + list(confounders)

    if analysis_type == "logistic_regression":
        params = {
            "outcome": outcome,
            "predictors": all_predictors
        }
    elif analysis_type == "cox_regression":
        params = {
            "time": request.get("time"),
            "event": request.get("event"),
            "covariates": all_predictors
        }
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported analysis type: {analysis_type}")

    # Run analysis
    result = analysis_service.run_analysis(session.data, analysis_type, params)

    if result.get("success"):
        session_manager.add_analysis(session_id, result)

    return JSONResponse(content=sanitize_for_json(result))


@router.post("/generate-report/{session_id}")
async def generate_results_report(session_id: str):
    """Generate a publication-ready Results section using LLM.

    Analyzes all completed analyses for the session and generates
    a formatted Results section suitable for a research paper.
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if not session.analyses:
        raise HTTPException(
            status_code=400,
            detail="No analyses found. Please run analyses before generating a report."
        )

    # Get research question from analysis plan
    research_question = ""
    if session.analysis_plan:
        research_question = session.analysis_plan.get("research_question", "")

    # Prepare data description
    data_description = {
        "n_rows": len(session.data) if session.data is not None else 0,
        "n_columns": len(session.data.columns) if session.data is not None else 0,
        "filename": session.filename or "Unknown",
        "variables": list(session.data.columns) if session.data is not None else []
    }

    # Generate report using LLM
    try:
        report = llm_service.generate_results_section(
            research_question=research_question,
            analyses=session.analyses,
            data_description=data_description
        )

        # Store the report in session
        session.results_report = report

        # Add to chat history
        session_manager.add_chat_message(
            session_id, "assistant",
            "Results section report generated successfully."
        )

        return JSONResponse(content={
            "success": True,
            "report": report,
            "research_question": research_question,
            "n_analyses": len(session.analyses)
        })

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate report: {str(e)}"
        )


@router.get("/report/{session_id}")
async def get_results_report(session_id: str):
    """Get the previously generated Results section report."""
    session = session_manager.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if not hasattr(session, 'results_report') or not session.results_report:
        return JSONResponse(content={
            "success": False,
            "message": "No report generated yet. Use POST /chat/generate-report/{session_id} to generate one."
        })

    return JSONResponse(content={
        "success": True,
        "report": session.results_report
    })


@router.post("/model-diagnostics/{session_id}")
async def run_model_diagnostics(session_id: str, request: Dict[str, Any]):
    """Run model diagnostics for regression models.

    Request body:
    {
        "model_type": "logistic" | "linear" | "cox",
        "outcome": "column_name",
        "predictors": ["var1", "var2"],
        "time": "time_col",  # For Cox only
        "event": "event_col"  # For Cox only
    }
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if session.data is None:
        raise HTTPException(status_code=400, detail="No data loaded")

    params = {
        "model_type": request.get("model_type", "logistic"),
        "outcome": request.get("outcome"),
        "predictors": request.get("predictors", []),
        "time": request.get("time"),
        "event": request.get("event")
    }

    result = analysis_service.run_analysis(session.data, "model_diagnostics", params)

    if result.get("success"):
        result["test_name"] = "Model Diagnostics"
        result["rationale"] = "Assess model assumptions and fit quality"
        session_manager.add_analysis(session_id, result)

    return JSONResponse(content=sanitize_for_json(result))


@router.post("/subgroup-analysis/{session_id}")
async def run_subgroup_analysis(session_id: str, request: Dict[str, Any]):
    """Run stratified/subgroup analysis.

    Request body:
    {
        "analysis_type": "logistic_regression" | "cox_regression",
        "stratify_by": "sex",
        "outcome": "mortality",
        "predictors": ["treatment"],
        "time": "time_col",  # For Cox only
        "event": "event_col"  # For Cox only
    }
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if session.data is None:
        raise HTTPException(status_code=400, detail="No data loaded")

    result = analysis_service.run_analysis(session.data, "subgroup_analysis", request)

    if result.get("success"):
        result["test_name"] = f"Subgroup Analysis by {request.get('stratify_by')}"
        result["rationale"] = "Examine treatment effects across different subgroups"
        session_manager.add_analysis(session_id, result)

    return JSONResponse(content=sanitize_for_json(result))


@router.get("/missing-data/{session_id}")
async def analyze_missing_data(session_id: str, variables: str = None):
    """Analyze missing data patterns.

    Query params:
    - variables: Comma-separated list of variables to analyze (optional)
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if session.data is None:
        raise HTTPException(status_code=400, detail="No data loaded")

    params = {}
    if variables:
        params["variables"] = [v.strip() for v in variables.split(",")]

    result = analysis_service.run_analysis(session.data, "missing_data_analysis", params)

    if result.get("success"):
        result["test_name"] = "Missing Data Analysis"
        result["rationale"] = "Understand patterns of missing data for handling strategy"

    return JSONResponse(content=sanitize_for_json(result))


@router.post("/propensity-score/{session_id}")
async def run_propensity_score_analysis(session_id: str, request: Dict[str, Any]):
    """Run propensity score analysis (matching or IPTW).

    Request body:
    {
        "treatment": "treatment_var",
        "outcome": "outcome_var",
        "covariates": ["age", "sex", ...],
        "method": "matching" | "iptw",
        "caliper": 0.2,  # For matching
        "outcome_type": "binary" | "continuous" | "survival"
    }
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if session.data is None:
        raise HTTPException(status_code=400, detail="No data loaded")

    result = analysis_service.run_analysis(session.data, "propensity_score_analysis", request)

    if result.get("success"):
        method = request.get("method", "matching")
        result["test_name"] = f"Propensity Score {method.upper()}"
        result["rationale"] = "Adjust for confounding using propensity scores"
        session_manager.add_analysis(session_id, result)

    return JSONResponse(content=sanitize_for_json(result))


@router.post("/power-calculation")
async def calculate_power(request: Dict[str, Any]):
    """Calculate power or sample size for study planning.

    Request body:
    {
        "test_type": "t_test" | "chi_square" | "survival",
        "effect_size": 0.5,  # Cohen's d, odds ratio, or hazard ratio
        "alpha": 0.05,
        "power": 0.8,  # Target power (for sample size calc)
        "n": 100,  # Sample size (for power calc)
        "ratio": 1,  # Allocation ratio
        "two_sided": true
    }
    """
    result = analysis_service.run_analysis(None, "power_calculation", request)

    if result.get("success"):
        result["test_name"] = "Power Analysis"
        result["rationale"] = "Determine required sample size or statistical power"

    return JSONResponse(content=sanitize_for_json(result))


@router.post("/export/{session_id}")
async def export_results(session_id: str, request: Dict[str, Any]):
    """Export analysis results in various formats.

    Request body:
    {
        "format": "word" | "latex" | "html" | "json",
        "include_tables": true,
        "include_figures": true,
        "analysis_ids": [1, 2, 3]  # Optional: specific analyses to export
    }
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if not session.analyses:
        raise HTTPException(status_code=400, detail="No analyses to export")

    export_format = request.get("format", "html")
    include_tables = request.get("include_tables", True)
    include_figures = request.get("include_figures", True)
    analysis_ids = request.get("analysis_ids")

    # Filter analyses if specific IDs provided
    if analysis_ids:
        analyses = [a for i, a in enumerate(session.analyses) if i in analysis_ids]
    else:
        analyses = session.analyses

    try:
        if export_format == "word":
            content = _export_to_word(analyses, session, include_tables, include_figures)
            return JSONResponse(content={
                "success": True,
                "format": "word",
                "filename": f"analysis_results_{session_id}.docx",
                "content_base64": content
            })
        elif export_format == "latex":
            content = _export_to_latex(analyses, include_tables, include_figures)
            return JSONResponse(content={
                "success": True,
                "format": "latex",
                "content": content
            })
        elif export_format == "html":
            content = _export_to_html(analyses, include_tables, include_figures)
            return JSONResponse(content={
                "success": True,
                "format": "html",
                "content": content
            })
        else:  # json
            return JSONResponse(content={
                "success": True,
                "format": "json",
                "analyses": sanitize_for_json(analyses)
            })
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


def _export_to_html(analyses: List[Dict], include_tables: bool, include_figures: bool) -> str:
    """Export analyses to HTML format."""
    html = """<!DOCTYPE html>
<html>
<head>
    <title>Statistical Analysis Results</title>
    <style>
        body { font-family: Arial, sans-serif; max-width: 900px; margin: 0 auto; padding: 20px; }
        h1 { color: #333; border-bottom: 2px solid #4f46e5; }
        h2 { color: #4f46e5; margin-top: 30px; }
        table { border-collapse: collapse; width: 100%; margin: 20px 0; }
        th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
        th { background-color: #4f46e5; color: white; }
        tr:nth-child(even) { background-color: #f9f9f9; }
        .interpretation { background-color: #f0f9ff; padding: 15px; border-left: 4px solid #4f46e5; margin: 15px 0; }
        .p-significant { color: #059669; font-weight: bold; }
        .p-not-significant { color: #6b7280; }
        img { max-width: 100%; margin: 15px 0; }
    </style>
</head>
<body>
    <h1>Statistical Analysis Results</h1>
"""

    for i, analysis in enumerate(analyses):
        test_name = analysis.get("test_name", f"Analysis {i+1}")
        rationale = analysis.get("rationale", "")
        results = analysis.get("results", {})
        interpretation = analysis.get("interpretation", "")

        html += f"<h2>{test_name}</h2>\n"
        if rationale:
            html += f"<p><em>{rationale}</em></p>\n"

        # Add key statistics
        if include_tables and results:
            html += "<table>\n<tr><th>Metric</th><th>Value</th></tr>\n"
            for key, value in results.items():
                if key not in ["coefficients", "table_data", "plot_base64", "forest_plot_base64", "km_plot_base64"]:
                    if isinstance(value, float):
                        html += f"<tr><td>{key}</td><td>{value:.4f}</td></tr>\n"
                    elif value is not None:
                        html += f"<tr><td>{key}</td><td>{value}</td></tr>\n"
            html += "</table>\n"

        # Add interpretation
        if interpretation:
            html += f'<div class="interpretation">{interpretation}</div>\n'

        # Add figures if present
        if include_figures:
            for key in ["plot_base64", "forest_plot_base64", "km_plot_base64"]:
                if key in results and results[key]:
                    html += f'<img src="data:image/png;base64,{results[key]}" alt="{key}">\n'

    html += "</body>\n</html>"
    return html


def _export_to_latex(analyses: List[Dict], include_tables: bool, include_figures: bool) -> str:
    """Export analyses to LaTeX format."""
    latex = r"""\documentclass{article}
\usepackage{booktabs}
\usepackage{graphicx}
\usepackage{float}

\begin{document}

\section{Statistical Analysis Results}

"""

    for i, analysis in enumerate(analyses):
        test_name = analysis.get("test_name", f"Analysis {i+1}")
        rationale = analysis.get("rationale", "")
        results = analysis.get("results", {})
        interpretation = analysis.get("interpretation", "")

        latex += f"\\subsection{{{test_name}}}\n\n"
        if rationale:
            latex += f"\\textit{{{rationale}}}\n\n"

        # Add key statistics as a table
        if include_tables and results:
            latex += "\\begin{table}[H]\n\\centering\n\\begin{tabular}{ll}\n\\toprule\n"
            latex += "Metric & Value \\\\\n\\midrule\n"
            for key, value in results.items():
                if key not in ["coefficients", "table_data", "plot_base64", "forest_plot_base64", "km_plot_base64"]:
                    if isinstance(value, float):
                        latex += f"{key.replace('_', ' ').title()} & {value:.4f} \\\\\n"
                    elif value is not None:
                        latex += f"{key.replace('_', ' ').title()} & {value} \\\\\n"
            latex += "\\bottomrule\n\\end{tabular}\n\\end{table}\n\n"

        # Add interpretation
        if interpretation:
            # Clean up markdown formatting for LaTeX
            clean_interp = interpretation.replace("**", "\\textbf{").replace("*", "")
            latex += f"\\textit{{{clean_interp}}}\n\n"

    latex += "\\end{document}"
    return latex


def _export_to_word(analyses: List[Dict], session, include_tables: bool, include_figures: bool) -> str:
    """Export analyses to Word format (returns base64 encoded docx)."""
    import base64
    from io import BytesIO

    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed. Run: pip install python-docx")

    doc = Document()

    # Title
    title = doc.add_heading("Statistical Analysis Results", 0)

    for i, analysis in enumerate(analyses):
        test_name = analysis.get("test_name", f"Analysis {i+1}")
        rationale = analysis.get("rationale", "")
        results = analysis.get("results", {})
        interpretation = analysis.get("interpretation", "")

        # Add section heading
        doc.add_heading(test_name, level=1)

        if rationale:
            p = doc.add_paragraph()
            p.add_run(rationale).italic = True

        # Add key statistics as a table
        if include_tables and results:
            table_data = [(k, v) for k, v in results.items()
                         if k not in ["coefficients", "table_data", "plot_base64", "forest_plot_base64", "km_plot_base64"]
                         and v is not None]

            if table_data:
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Metric'
                hdr_cells[1].text = 'Value'

                for key, value in table_data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = key.replace('_', ' ').title()
                    if isinstance(value, float):
                        row_cells[1].text = f"{value:.4f}"
                    else:
                        row_cells[1].text = str(value)

        # Add interpretation
        if interpretation:
            doc.add_paragraph()
            # Simple markdown parsing
            for line in interpretation.split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line.replace('**', '').replace('*', ''))

        # Add figures if present
        if include_figures:
            for key in ["plot_base64", "forest_plot_base64", "km_plot_base64"]:
                if key in results and results[key]:
                    try:
                        img_data = base64.b64decode(results[key])
                        img_stream = BytesIO(img_data)
                        doc.add_picture(img_stream, width=Inches(5))
                    except Exception:
                        pass

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return base64.b64encode(buffer.read()).decode('utf-8')


@router.websocket("/ws/{session_id}")
async def websocket_chat(websocket: WebSocket, session_id: str):
    """WebSocket endpoint for real-time chat."""
    await websocket.accept()

    session = session_manager.get_session(session_id)
    if not session:
        await websocket.send_json({"error": "Session not found"})
        await websocket.close()
        return

    try:
        while True:
            data = await websocket.receive_json()
            message = data.get("message", "")

            # Build context
            context = {
                "data_loaded": session.data is not None,
                "filename": session.filename,
                "n_rows": len(session.data) if session.data is not None else 0,
                "variables": list(session.data.columns)[:20] if session.data is not None else []
            }

            # Get response
            response = llm_service.chat(message, context, session.chat_history)

            # Store messages
            session_manager.add_chat_message(session_id, "user", message)
            session_manager.add_chat_message(session_id, "assistant", response["message"])

            # Send response
            await websocket.send_json({
                "type": "message",
                "content": response["message"],
                "action": response.get("action")
            })

    except WebSocketDisconnect:
        pass
