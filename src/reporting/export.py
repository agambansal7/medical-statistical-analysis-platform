"""Report Export Module.

Export reports to various formats:
- Microsoft Word (.docx)
- LaTeX (.tex)
- PDF
- HTML
- Markdown
"""

from typing import Dict, Any, List, Optional
from pathlib import Path
from datetime import datetime
from .report_generator import FullReport, ReportSection


class ReportExporter:
    """Export reports to various formats."""

    def __init__(self, output_dir: str = "./reports"):
        """Initialize exporter.

        Args:
            output_dir: Directory for exported files
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def export(
        self,
        report: FullReport,
        format: str = "docx",
        filename: Optional[str] = None
    ) -> str:
        """Export report to specified format.

        Args:
            report: Report to export
            format: Output format (docx, latex, pdf, html, markdown)
            filename: Output filename (auto-generated if None)

        Returns:
            Path to exported file
        """
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"report_{timestamp}"

        if format == "docx":
            return self._export_docx(report, filename)
        elif format == "latex":
            return self._export_latex(report, filename)
        elif format == "html":
            return self._export_html(report, filename)
        elif format == "markdown":
            return self._export_markdown(report, filename)
        else:
            raise ValueError(f"Unsupported format: {format}")

    def _export_docx(self, report: FullReport, filename: str) -> str:
        """Export to Word document."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx is required for Word export")

        doc = Document()

        # Title
        title = doc.add_heading(report.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Sections
        for section_name, section in report.sections.items():
            self._add_section_to_docx(doc, section)

        # Tables
        for table in report.tables:
            self._add_table_to_docx(doc, table)

        # References
        if report.references:
            doc.add_heading("References", level=1)
            for i, ref in enumerate(report.references, 1):
                doc.add_paragraph(f"{i}. {ref}")

        # Save
        output_path = self.output_dir / f"{filename}.docx"
        doc.save(str(output_path))

        return str(output_path)

    def _add_section_to_docx(self, doc, section: ReportSection):
        """Add section to Word document."""
        from docx.shared import Pt

        doc.add_heading(section.title, level=section.level)

        # Add content paragraphs
        for paragraph in section.content.split('\n\n'):
            p = doc.add_paragraph(paragraph.strip())
            p.style.font.size = Pt(11)

        # Add subsections
        for subsection in section.subsections:
            self._add_section_to_docx(doc, subsection)

    def _add_table_to_docx(self, doc, table: Dict[str, Any]):
        """Add table to Word document."""
        from docx.shared import Inches

        # Add table title
        doc.add_paragraph(f"Table: {table.get('title', 'Untitled')}")

        # Create table
        data = table.get('data', [])
        columns = table.get('columns', [])

        if not data or not columns:
            return

        word_table = doc.add_table(rows=1, cols=len(columns))
        word_table.style = 'Table Grid'

        # Header row
        header_cells = word_table.rows[0].cells
        for i, col in enumerate(columns):
            header_cells[i].text = str(col)

        # Data rows
        for row_data in data:
            row_cells = word_table.add_row().cells
            for i, col in enumerate(columns):
                row_cells[i].text = str(row_data.get(col, ''))

        # Footnotes
        for footnote in table.get('footnotes', []):
            doc.add_paragraph(footnote, style='Caption')

    def _export_latex(self, report: FullReport, filename: str) -> str:
        """Export to LaTeX document."""
        lines = [
            r"\documentclass{article}",
            r"\usepackage{booktabs}",
            r"\usepackage{graphicx}",
            r"\usepackage{hyperref}",
            "",
            r"\begin{document}",
            "",
            r"\title{" + self._escape_latex(report.title) + "}",
            r"\maketitle",
            ""
        ]

        # Sections
        for section_name, section in report.sections.items():
            lines.extend(self._section_to_latex(section))

        # Tables
        for table in report.tables:
            lines.extend(self._table_to_latex(table))

        # References
        if report.references:
            lines.extend([
                "",
                r"\section{References}",
                r"\begin{enumerate}"
            ])
            for ref in report.references:
                lines.append(r"\item " + self._escape_latex(ref))
            lines.append(r"\end{enumerate}")

        lines.append(r"\end{document}")

        # Save
        output_path = self.output_dir / f"{filename}.tex"
        with open(output_path, 'w') as f:
            f.write('\n'.join(lines))

        return str(output_path)

    def _section_to_latex(self, section: ReportSection) -> List[str]:
        """Convert section to LaTeX."""
        section_cmds = ['section', 'subsection', 'subsubsection']
        cmd = section_cmds[min(section.level - 1, 2)]

        lines = [
            "",
            f"\\{cmd}{{{self._escape_latex(section.title)}}}",
            "",
            self._escape_latex(section.content),
        ]

        for subsection in section.subsections:
            lines.extend(self._section_to_latex(subsection))

        return lines

    def _table_to_latex(self, table: Dict[str, Any]) -> List[str]:
        """Convert table to LaTeX."""
        data = table.get('data', [])
        columns = table.get('columns', [])

        if not data or not columns:
            return []

        col_spec = '|' + 'l|' * len(columns)

        lines = [
            "",
            r"\begin{table}[htbp]",
            r"\centering",
            r"\caption{" + self._escape_latex(table.get('title', '')) + "}",
            r"\begin{tabular}{" + col_spec + "}",
            r"\toprule",
            " & ".join(self._escape_latex(c) for c in columns) + r" \\",
            r"\midrule"
        ]

        for row_data in data:
            row_values = [str(row_data.get(col, '')) for col in columns]
            lines.append(" & ".join(self._escape_latex(v) for v in row_values) + r" \\")

        lines.extend([
            r"\bottomrule",
            r"\end{tabular}",
            r"\end{table}"
        ])

        return lines

    def _export_html(self, report: FullReport, filename: str) -> str:
        """Export to HTML document."""
        html_parts = [
            "<!DOCTYPE html>",
            "<html>",
            "<head>",
            f"<title>{report.title}</title>",
            "<style>",
            "body { font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }",
            "table { border-collapse: collapse; width: 100%; margin: 20px 0; }",
            "th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }",
            "th { background-color: #f2f2f2; }",
            "h1 { text-align: center; }",
            "</style>",
            "</head>",
            "<body>",
            f"<h1>{report.title}</h1>"
        ]

        # Sections
        for section_name, section in report.sections.items():
            html_parts.append(self._section_to_html(section))

        # Tables
        for table in report.tables:
            html_parts.append(self._table_to_html(table))

        # References
        if report.references:
            html_parts.append("<h2>References</h2>")
            html_parts.append("<ol>")
            for ref in report.references:
                html_parts.append(f"<li>{ref}</li>")
            html_parts.append("</ol>")

        html_parts.extend(["</body>", "</html>"])

        # Save
        output_path = self.output_dir / f"{filename}.html"
        with open(output_path, 'w') as f:
            f.write('\n'.join(html_parts))

        return str(output_path)

    def _section_to_html(self, section: ReportSection) -> str:
        """Convert section to HTML."""
        level = min(section.level + 1, 6)  # h2 to h6
        html = f"<h{level}>{section.title}</h{level}>\n"

        for paragraph in section.content.split('\n\n'):
            html += f"<p>{paragraph.strip()}</p>\n"

        for subsection in section.subsections:
            html += self._section_to_html(subsection)

        return html

    def _table_to_html(self, table: Dict[str, Any]) -> str:
        """Convert table to HTML."""
        data = table.get('data', [])
        columns = table.get('columns', [])

        if not data or not columns:
            return ""

        html = f"<h4>Table: {table.get('title', '')}</h4>\n<table>\n<thead>\n<tr>\n"

        for col in columns:
            html += f"<th>{col}</th>\n"
        html += "</tr>\n</thead>\n<tbody>\n"

        for row_data in data:
            html += "<tr>\n"
            for col in columns:
                html += f"<td>{row_data.get(col, '')}</td>\n"
            html += "</tr>\n"

        html += "</tbody>\n</table>\n"

        return html

    def _export_markdown(self, report: FullReport, filename: str) -> str:
        """Export to Markdown document."""
        lines = [f"# {report.title}", ""]

        # Sections
        for section_name, section in report.sections.items():
            lines.extend(self._section_to_markdown(section))

        # Tables
        for table in report.tables:
            lines.extend(self._table_to_markdown(table))

        # References
        if report.references:
            lines.extend(["", "## References", ""])
            for i, ref in enumerate(report.references, 1):
                lines.append(f"{i}. {ref}")

        # Save
        output_path = self.output_dir / f"{filename}.md"
        with open(output_path, 'w') as f:
            f.write('\n'.join(lines))

        return str(output_path)

    def _section_to_markdown(self, section: ReportSection) -> List[str]:
        """Convert section to Markdown."""
        header = '#' * (section.level + 1)
        lines = [
            f"{header} {section.title}",
            "",
            section.content,
            ""
        ]

        for subsection in section.subsections:
            lines.extend(self._section_to_markdown(subsection))

        return lines

    def _table_to_markdown(self, table: Dict[str, Any]) -> List[str]:
        """Convert table to Markdown."""
        data = table.get('data', [])
        columns = table.get('columns', [])

        if not data or not columns:
            return []

        lines = [
            f"**Table: {table.get('title', '')}**",
            "",
            "| " + " | ".join(columns) + " |",
            "| " + " | ".join(['---'] * len(columns)) + " |"
        ]

        for row_data in data:
            values = [str(row_data.get(col, '')) for col in columns]
            lines.append("| " + " | ".join(values) + " |")

        lines.append("")

        return lines

    def _escape_latex(self, text: str) -> str:
        """Escape special LaTeX characters."""
        replacements = {
            '&': r'\&',
            '%': r'\%',
            '$': r'\$',
            '#': r'\#',
            '_': r'\_',
            '{': r'\{',
            '}': r'\}',
            '~': r'\textasciitilde{}',
            '^': r'\^{}',
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        return text
